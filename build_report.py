"""Génère le rapport DOCX et un aperçu statique de la carte."""
from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
DATA = json.loads((ROOT / "data_source.json").read_text(encoding="utf-8"))["users"]
BLUE = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)


def create_map_preview():
    fig, ax = plt.subplots(figsize=(10, 4.4))
    ax.set_facecolor("#EAF4F8")
    ax.set_xlim(-180, 180); ax.set_ylim(-90, 90)
    ax.set_xticks(range(-180, 181, 60)); ax.set_yticks(range(-90, 91, 30))
    ax.grid(color="white", linewidth=1.2)
    for u in DATA:
        lat = float(u["address"]["geo"]["lat"]); lon = float(u["address"]["geo"]["lng"])
        ax.scatter(lon, lat, s=70, color="#D35400", edgecolor="white", linewidth=1.2, zorder=3)
        ax.annotate(str(u["id"]), (lon, lat), ha="center", va="center", fontsize=7, color="white", weight="bold")
    ax.set(title="Aperçu des coordonnées utilisées par la carte Folium", xlabel="Longitude", ylabel="Latitude")
    fig.tight_layout(); fig.savefig(ROOT / "map_preview.png", dpi=180, bbox_inches="tight"); plt.close(fig)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(7)
    for r in p.runs: r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after in (("Heading 1",16,16,8),("Heading 2",13,12,6),("Heading 3",12,8,4)):
        st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = BLUE
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

    header = sec.header.paragraphs[0]
    header.text = "COLLECTE ET EXPLORATION DES DONNÉES  |  MINI-PROJET"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs: r.font.size = Pt(8); r.font.color.rgb = MUTED
    page_number(sec.footer.paragraphs[0])

    # Page 1
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(5)
    r = p.add_run("RAPPORT DE MINI-PROJET"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = BLUE
    p = doc.add_paragraph("Pipeline de données des utilisateurs JSONPlaceholder")
    p.paragraph_format.space_after = Pt(18)
    for r in p.runs: r.font.size = Pt(15); r.font.color.rgb = MUTED
    p = doc.add_paragraph("Étudiant(e) : [Nom et prénom]     |     Année universitaire : 2025–2026")
    for r in p.runs: r.bold = True; r.font.size = Pt(9.5)

    add_heading(doc, "1. Présentation du travail", 1)
    doc.add_paragraph(
        "Ce mini-projet met en œuvre un pipeline complet de collecte, stockage et exploration. "
        "Le script Python interroge l’API JSONPlaceholder, extrait les dix utilisateurs et charge "
        "leurs informations dans une base PostgreSQL. Le notebook Jupyter exécute ensuite les "
        "requêtes SQL demandées, analyse les extensions d’e-mail et construit une carte interactive."
    )
    add_heading(doc, "2. Pipeline et modèle de données", 1)
    doc.add_paragraph(
        "Le chargement utilise un upsert PostgreSQL : une seconde exécution met à jour les lignes "
        "existantes au lieu de créer des doublons. La table utilisateurs comporte id (clé primaire), "
        "nom, email et ville. Les colonnes latitude et longitude ont été ajoutées pour conserver les "
        "coordonnées address.geo nécessaires à la carte."
    )
    add_heading(doc, "Étapes réalisées", 2)
    for text in (
        "Extraction HTTP des données JSON et validation du nombre de lignes.",
        "Transformation des champs imbriqués address.city et address.geo.",
        "Création de la table puis insertion idempotente des dix enregistrements.",
        "Connexion du notebook à PostgreSQL, analyses SQL et visualisations.",
    ):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()
    # Page 2
    add_heading(doc, "3. Vérification de la base PostgreSQL", 1)
    doc.add_paragraph(
        "La capture ci-dessous présente les quatre champs demandés pour les dix lignes enregistrées. "
        "Elle est produite directement depuis une requête SELECT du notebook."
    )
    doc.add_picture(str(ROOT / "table_sql.png"), width=Inches(6.45))
    add_caption(doc, "Figure 1 — Table utilisateurs remplie dans PostgreSQL.")
    add_heading(doc, "4. Résultats des requêtes SQL", 1)
    doc.add_paragraph("La requête COUNT(*) retourne 10 utilisateurs. Le filtrage par suffixe retourne quatre lignes :")
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False; widths = [Inches(2.0), Inches(3.1), Inches(1.4)]
    headers = ["Nom", "E-mail", "Suffixe"]
    for i, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = width; cell.text = text; set_cell_margins(cell)
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd")); cell._tc.tcPr[-1].set(qn("w:fill"), "D9EAF7")
        for r in cell.paragraphs[0].runs: r.bold = True
    selected = [u for u in DATA if u["email"].lower().endswith((".net", ".biz"))]
    for u in selected:
        row = table.add_row().cells
        vals = [u["name"], u["email"], "." + u["email"].split(".")[-1].lower()]
        for cell, width, val in zip(row, widths, vals): cell.width = width; cell.text = val; set_cell_margins(cell)
    add_caption(doc, "Tableau 1 — Utilisateurs dont l’e-mail se termine par .net ou .biz.")

    doc.add_page_break()
    # Page 3
    add_heading(doc, "5. Visualisations et interprétation", 1)
    doc.add_picture(str(ROOT / "extensions_email.png"), width=Inches(5.75))
    add_caption(doc, "Figure 2 — Répartition des utilisateurs par extension d’adresse e-mail.")
    p = doc.add_paragraph()
    p.add_run("Lecture : ").bold = True
    p.add_run("l’extension .biz domine avec 3 utilisateurs. Les sept autres extensions n’apparaissent qu’une fois chacune ; .net compte 1 utilisateur.")
    doc.add_picture(str(ROOT / "map_preview.png"), width=Inches(5.85))
    add_caption(doc, "Figure 3 — Aperçu statique des 10 marqueurs de la carte Folium interactive.")
    p = doc.add_paragraph()
    p.add_run("Lecture : ").bold = True
    p.add_run(
        "les coordonnées de démonstration sont très dispersées. Dans map_utilisateurs.html et dans le notebook, "
        "chaque marqueur affiche le nom et la ville au survol ou au clic. Ces positions fictives ne représentent pas des domiciles réels."
    )

    doc.core_properties.title = "Mini-projet — Collecte et exploration des données"
    doc.core_properties.subject = "Pipeline JSONPlaceholder vers PostgreSQL"
    doc.core_properties.author = "[Nom et prénom]"
    doc.save(ROOT / "rapport.docx")


if __name__ == "__main__":
    create_map_preview()
    build_docx()
    print("rapport.docx et map_preview.png créés")
